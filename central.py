import PySimpleGUI as sg 
import base64
import tkinter as tk
import ctypes
import mysql.connector
import docx
import os
import time
import threading


sg.change_look_and_feel('DarkGreen')
ip = sg.popup_get_text('Digite o ip do server:', 'Ex: 192.168.0.100')
senha = sg.popup_get_text('Digite a senha de acesso:', 'senha', password_char = "*",)




dbestoque = mysql.connector.connect(
        host= ip,
        user = 'igaoadm',
        passwd = senha,
        database='estoque'
)

cursor = dbestoque.cursor()

cursor.execute('''CREATE DATABASE IF NOT EXISTS estoque''')


cursor.execute(''' CREATE TABLE IF NOT EXISTS produtos(
                    Nome MEDIUMTEXT NOT NULL,
                    Quantidade INTEGER NOT NULL,
                    Preco FLOAT NOT NULL,
                    Id TEXT NOT NULL,
                    Classe TEXT NOT NULL)''')

cursor.execute(''' CREATE TABLE IF NOT EXISTS categorias(
                    Categoria TEXT NOT NULL)''')   
                    

cursor.execute(''' CREATE TABLE IF NOT EXISTS televisao(
                    Tvs LONGTEXT NOT NULL)''')






def escreve(Nome, Quantidade, Preco, Id, Classe):
    cursor.execute(''' INSERT INTO produtos (Nome, Quantidade, Preco, Id, Classe) VALUES(%s, %s, %s, %s, %s)''', [Nome, Quantidade, Preco, Id, Classe])
    dbestoque.commit()

def escreve_TV(Tvs):
    cursor.execute(''' INSERT INTO televisao (Tvs) VALUES(%s)''', [Tvs])
    dbestoque.commit()


def escreve_class(Categoria):
    cursor.execute(''' INSERT INTO categorias (Categoria) VALUES(%s)''', [Categoria])
    dbestoque.commit()

def escreve_marca(Marca):
    cursor.execute(''' INSERT INTO Marcas_db (Marca) VALUES(%s)''', [Marca])
    dbestoque.commit()


def delete(x):
    cursor.execute('''delete from produtos where Nome = %s''', (x,))
    dbestoque.commit

def delete_tv(t):
    cursor.execute('''delete from televisao where Tvs LIKE %s''',('%'+t+'%',))


def read_task0():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT * from produtos''')
    data = cursor.fetchall()
    dbestoque.commit()
    return data
def read_task():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Nome from produtos''')
    data = cursor.fetchall()
    dbestoque.commit()
    return data
def read_task1():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Quantidade from produtos''')
    data1 = cursor.fetchall()
    dbestoque.commit()
    return data1
def read_task2():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Preco from produtos''')
    data2 = cursor.fetchall()
    dbestoque.commit()
    return data2
def read_task3():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Id from produtos''')
    data3 = cursor.fetchall()
    dbestoque.commit()
    return data3

def read_task4():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Classe from produtos''')
    data3 = cursor.fetchall()
    dbestoque.commit()
    return data3

def read_task5():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Categoria from categorias''')
    data4 = cursor.fetchall()
    dbestoque.commit()
    return data4
######################################################READ TASK TV##########################################
def TV_readTask():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT * from televisao''')
    data = cursor.fetchall()
    dbestoque.commit()
    return data

def marcas_task():
    cursor = dbestoque.cursor()
    cursor.execute('''SELECT Marca from Marcas_db''')
    data = cursor.fetchall()
    dbestoque.commit()
    return data




############################################################################################################

def baixaTV(x,z):
    cursor.execute('''update televisao set Tvs=%s where Tvs= %s''', (x,z))
    dbestoque.commit()

def vender(q, h):
    cursor.execute('''update produtos set Quantidade =%s where Nome= %s''', (q,h,))
    dbestoque.commit()

def add_estoque(add_q, h):
    cursor.execute('''update produtos set Quantidade =%s where Nome= %s''', (add_q,h,))
    dbestoque.commit()

def Editar(p, h):
    cursor.execute('''update produtos set Preco =%s where Nome=%s''', (p, h))
    dbestoque.commit()

def EditarID(p, h):
    cursor.execute('''update produtos set Id =%s where Nome=%s''', (p, h))
    dbestoque.commit()

def classificacao(w, h):
    cursor.execute('''update produtos set Classe =%s where Nome=%s''', (w, h))
    dbestoque.commit()



def busca0(c):
    cursor.execute('''SELECT Id from produtos where Nome LIKE %s''',('%'+c+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado

def busca1(c):
    cursor.execute('''SELECT Nome from produtos WHERE Nome LIKE %s ''',('%'+c+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado

def busca2(c):
    cursor.execute('''SELECT Quantidade from produtos where Nome LIKE %s''',('%'+c+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado

def busca2_teste(y):
    cursor.execute('''SELECT Quantidade from produtos where Nome=%s''',(y,))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado

def busca3(c):
    cursor.execute('''SELECT Preco from produtos where Nome LIKE %s''',('%'+c+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def busca4(c):
    cursor.execute('''SELECT Classe from produtos where Nome LIKE %s''',('%'+c+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado

def filtrar(f):
    cursor.execute('''SELECT Nome from produtos where Classe= %s && Quantidade <= 2''',(f,))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def filtrar2(f):
    cursor.execute('''SELECT Id from produtos where Classe= %s and Quantidade <= 2''',(f,))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def filtrar3(f):
    cursor.execute('''SELECT Quantidade from produtos where Classe= %s and Quantidade <= 2''',(f,))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def filtrar4(f):
    cursor.execute('''SELECT Preco from produtos where Classe= %s and Quantidade <= 2''',(f,))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado


def filtrar_com_marca(Marca):
    cursor.execute('''SELECT Nome from produtos where Nome LIKE %s and Quantidade <= 2  ''',('%'+Marca+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def filtrar2_com_marca(Marca):
    cursor.execute('''SELECT Id from produtos where Nome LIKE %s and Quantidade <= 2  ''',('%'+Marca+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def filtrar3_com_marca(Marca):
    cursor.execute('''SELECT Quantidade from produtos where Nome LIKE %s and Quantidade <= 2  ''',('%'+Marca+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado
def filtrar4_com_marca(Marca):
    cursor.execute('''SELECT Preco from produtos where Nome LIKE %s and Quantidade <= 2  ''',('%'+Marca+'%',))
    resultado = cursor.fetchall()
    dbestoque.commit()
    return resultado



    
##############################################main############################################

#########################################ADICIONAR TV ########################################



def AdicionarTelevisao():
    with open("logo.ico", "rb") as f:
        my_icon = base64.b64encode(f.read())
        sg.set_options(icon=my_icon)


    #mudar tema
    sg.change_look_and_feel('DarkGreen')

    marcas = marcas_task()
    Tvs = TV_readTask()
    treedata = sg.TreeData()
    lista = []
    for a in Tvs:
        lista.append(a[0])


    for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
            separated = val.split(".") # creates list of dot-pathed names
            fullname = '' 
            for f in separated:
                parent=fullname
                if fullname =='': # first time it runs... or any time I read a new entry starting at root
                    fullname=f
                    if i == 0: # only insert root on first run
                        treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                else:
                    fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                    if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                        treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
    
            
        
    


    #layout
    
    layout =[
            [sg.Text('Selecione a Marca da TV:',size=(17,0)),sg.InputCombo(marcas,size=(25,0),key='marca_TV')],
            [sg.Text('Digite o Modelo da TV:',size=(17,0)),sg.Input(do_not_clear=False, size=(25,0),key='modelo_TV')],
            [sg.Text('Digite a quantidade de Principal caso tenha:',size=(17,0)),sg.Input(do_not_clear=False,size=(25,0),key='principal_TV')],
            [sg.Text('Digite a quantidade de Fonte caso tenha',size=(17,0)),sg.Input(do_not_clear=False,size=(25,0),key='fonte_TV')],
            [sg.Button('Adicionar')], 
            [ sg.Button('File') ],
            [ sg.Tree(data=treedata, headings=['col1', 'col2', 'col3'], auto_size_columns=True, num_rows=10, col0_width=30, key='_TREE_', show_expanded=False,)],
            [sg.Button('Deletar', font=(10)), sg.Button('Realizar venda',pad=((10,180),12)), sg.Button('Adicionar placa')],
            [sg.Button('Sair',size=(10,2)),sg.Button('⬅ Voltar',font=(10))]
            ]
    #janela
    window = sg.Window("adicionar TV ao Estoque",layout)

    while True:
        event,values = window.read()

      
            
        if event == 'Adicionar':
            try: 
                MarcaIntra = values['marca_TV']
                Marca = MarcaIntra[0]
                Modelo = values['modelo_TV']
                PrincipalIntra = values['principal_TV']
                FonteIntra = values['fonte_TV']
                Principal = ("Principal=%s"%(PrincipalIntra))
                Fonte = ("Fonte=%s"%(FonteIntra))

                if Marca != '':
                    placas = [Principal,Fonte]
                    for i in placas:
                        Tvs = ("Tvs.%s.%s.%s"%(Marca,Modelo,i))
                        escreve_TV(Tvs)

                Tvs = TV_readTask()
                treedata = sg.TreeData()
                lista = []
                for a in Tvs:
                    lista.append(a[0])
                
                icone= b'./icone.png'
                for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
                    separated = val.split(".") # creates list of dot-pathed names
                    fullname = '' 
                    for f in separated:
                        parent=fullname
                        if fullname =='': # first time it runs... or any time I read a new entry starting at root
                            fullname=f
                            if i == 0: # only insert root on first run
                                treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                        else:
                            fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                            if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                                treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values

                window['_TREE_'].Update(values=treedata,key=fullname, value=None,text=f)
                window.refresh()                            

            except ValueError as erro:
                print("Coloque os valores corretamente") 

            
        if event == 'Deletar':
                teste = sg.popup_ok_cancel('Deseja Mesmo Deletar este item?')
                if teste == 'Cancel':
                    print('okay')
                else:
                    y = values['_TREE_'][0]
                    separated = y.split(".")
                    t = separated[2]
                    delete_tv(t)
                    
                Tvs = TV_readTask()
                treedata = sg.TreeData()
                lista = []
                for a in Tvs:
                    lista.append(a[0])
                
                icone= b'./icone.png'
                for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
                    separated = val.split(".") # creates list of dot-pathed names
                    fullname = '' 
                    for f in separated:
                        parent=fullname
                        if fullname =='': # first time it runs... or any time I read a new entry starting at root
                            fullname=f
                            if i == 0: # only insert root on first run
                                treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                        else:
                            fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                            if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                                treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values

                window['_TREE_'].Update(values=treedata,key=fullname, value=None,text=f)
                window.refresh()   

        if event == 'Realizar venda':
            placa = values['_TREE_']
            z = placa
            for i in placa:
                valores = i.split('.')
                Marca = valores[1]
                Modelo = valores[2]
                Placa = valores[3]

            
                if "Principal" in Placa:
                    amount = int(Placa[10])
                    result = str(amount - 1)
                    Placa = "Principal="+result
                    x = ("Tvs.%s.%s.%s"%(Marca,Modelo,Placa))
                    z = values['_TREE_'][0]
                    baixaTV(x,z)
                    
                    Tvs = TV_readTask()
                    treedata = sg.TreeData()
                    lista = []
                    for a in Tvs:
                        lista.append(a[0])
                    
                    for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
                        separated = val.split(".") # creates list of dot-pathed names
                        fullname = '' 
                        for f in separated:
                            parent=fullname
                            if fullname =='': # first time it runs... or any time I read a new entry starting at root
                                fullname=f
                                if i == 0: # only insert root on first run
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                            else:
                                fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                                if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                                    

                    window['_TREE_'].Update(values=treedata,key=fullname, value=None,text=f)
                    window.refresh()   

                elif "Fonte" in Placa:
                    amount = int(Placa[6])
                    result = str(amount - 1)
                    Placa = "Fonte="+result
                    x = ("Tvs.%s.%s.%s"%(Marca,Modelo,Placa))
                    z = values['_TREE_'][0]
                    baixaTV(x,z)

                    Tvs = TV_readTask()
                    treedata = sg.TreeData()
                    lista = []
                    for a in Tvs:
                        lista.append(a[0])

                    for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
                        separated = val.split(".") # creates list of dot-pathed names
                        fullname = '' 
                        for f in separated:
                            parent=fullname
                            if fullname =='': # first time it runs... or any time I read a new entry starting at root
                                fullname=f
                                if i == 0: # only insert root on first run
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                            else:
                                fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                                if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                                    

                    window['_TREE_'].Update(values=treedata,key=fullname, value=None,text=f)
                    window.refresh()


        if event == 'Adicionar placa':
            placa = values['_TREE_']
            z = placa
            for i in placa:
                valores = i.split('.')
                Marca = valores[1]
                Modelo = valores[2]
                Placa = valores[3]

            
                if "Principal" in Placa:
                    amount = int(Placa[10])
                    result = str(amount + 1)
                    Placa = "Principal="+result
                    x = ("Tvs.%s.%s.%s"%(Marca,Modelo,Placa))
                    z = values['_TREE_'][0]
                    baixaTV(x,z)
                    
                    Tvs = TV_readTask()
                    treedata = sg.TreeData()
                    lista = []
                    for a in Tvs:
                        lista.append(a[0])
                    
                    for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
                        separated = val.split(".") # creates list of dot-pathed names
                        fullname = '' 
                        for f in separated:
                            parent=fullname
                            if fullname =='': # first time it runs... or any time I read a new entry starting at root
                                fullname=f
                                if i == 0: # only insert root on first run
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                            else:
                                fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                                if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                                    

                    window['_TREE_'].Update(values=treedata,key=fullname, value=None,text=f)
                    
                    window.refresh()   

                elif "Fonte" in Placa:
                    amount = int(Placa[6])
                    result = str(amount + 1)
                    Placa = "Fonte="+result
                    x = ("Tvs.%s.%s.%s"%(Marca,Modelo,Placa))
                    z = values['_TREE_'][0]
                    baixaTV(x,z)

                    Tvs = TV_readTask()
                    treedata = sg.TreeData()
                    lista = []
                    for a in Tvs:
                        lista.append(a[0])

                    for i, val in enumerate(lista) :     #use enumerate and two values so I can keep track of index (non-standard in python for loops)
                        separated = val.split(".") # creates list of dot-pathed names
                        fullname = '' 
                        for f in separated:
                            parent=fullname
                            if fullname =='': # first time it runs... or any time I read a new entry starting at root
                                fullname=f
                                if i == 0: # only insert root on first run
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                            else:
                                fullname = fullname+'.'+f #rebuild the dot-pathed name step by step just like in the file tree example.
                                if fullname not in treedata.tree_dict: # Only create this key if it doesn't already exist                                     
                                    treedata.insert(parent, fullname, f, values=[] ); #enter first entry as the root module with no values
                                    

                    window['_TREE_'].Update(values=treedata,key=fullname, value=None,text=f)
                    window.refresh()  

        if event == "File":
            window['_TREE_'].treedata

        if event == sg.WIN_CLOSED or event == 'Sair':
            window.close()
            break

        if event == '⬅ Voltar':
            window.close()
            initi()
          
            
##############################################################################################

def AdicionarItem():
    with open("logo.ico", "rb") as f:
        my_icon = base64.b64encode(f.read())
        sg.set_options(icon=my_icon)



    #mudar tema
    sg.change_look_and_feel('DarkGreen')


    Nome = read_task()
    Quantidade = read_task1()
    Preco = read_task2()
    Id = read_task3()
    Classe = read_task4()
    Categoria = read_task5()

    cols = 5
    rows = len(Nome)
    col_width = 30
    
    
    all_listbox = [[sg.Listbox(Id, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {0}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(50, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {1}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Quantidade, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {2}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Preco, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {3}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE),sg.Listbox(Classe, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {4}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE)]]

    #layout
    layout = [
        [sg.Text('Digite o nome do item:',size=(17,0)),sg.Input(do_not_clear=False, size=(25,0),key='add_item',pad=((5,330),12)), sg.Button('➕Adicionar classe➕')],
        [sg.Text('A quantidade:',size=(17,0)),sg.Input(do_not_clear=False, size=(25,0),key='add_quantidade')],
        [sg.Text('Digite o preço do item:',size=(17,0)),sg.Input(do_not_clear=False,size=(25,0),key='add_preco')],
        [sg.Text('Digite o codigo de ID:',size=(17,0)),sg.Input(do_not_clear=False,size=(25,0),key='Cod_id')],
        [sg.Text('Selecione a categoria:',size=(17,0)),sg.InputCombo(Categoria,size=(25,0),key='combo')],
        [sg.Button('Adicionar')], 
        [sg.Text('ID',pad=((12, 200),12),font=('helvetica',12)), sg.Text('Produto',pad=((12, 195),12),font=('helvetica',12)) ,sg.Text('Quantidade',pad=((12, 35),12),font=('helvetica',12)), sg.Text('Preço',pad=((12, 40),12),font=('helvetica',12)), sg.Text('Classe',font=('helvetica',12))],
        [sg.Column(all_listbox, size=(811, 250), pad=(0, 0),scrollable=True,
         vertical_scroll_only=True, key='coluna')],
        [sg.Button('Deletar', font=(10))],
        [sg.Button('Sair',size=(10,2)),sg.Button('⬅ Voltar',font=(10))]
        ]
    #janela
    window = sg.Window("adicionar ao Estoque",layout)

    while True:
        event,values = window.read()
        try:
            if event.startswith('listbox'):
                        row = window[event].get_indexes()[0]
                        user_event = False
                        for i in range(cols):
                            listbox = window[f'listbox {i}']
                            listbox.set_value([])
                            listbox.Widget.selection_set(row)   # tkinter code
        except AttributeError:
            window.close()
            
        if event == 'Adicionar':
            try: 
                Nome = values['add_item']
                quantidadeIntra = values['add_quantidade']
                precoIntra = values['add_preco']
                Id = values['Cod_id']
                classe_nao_tratada = values['combo']
                Classe = str(classe_nao_tratada[0])
                Quantidade = int(quantidadeIntra)
                Preco = float(precoIntra)

                if Nome != '':
                    escreve(Nome, Quantidade, Preco, Id, Classe)
                Nome = read_task()
                rows = len(Nome)
                Quantidade = read_task1()
                Preco = read_task2()
                Id = read_task3()
                Classe = read_task4()
                for i in range(cols):
                    listbox = window[f'listbox {i}']
                    listbox.Widget.configure(height=rows)  
                window.find_element(f'listbox {0}').Update(Id)
                window.find_element(f'listbox {1}').Update(Nome)
                window.find_element(f'listbox {2}').Update(Quantidade)
                window.find_element(f'listbox {3}').Update(Preco)
                window.find_element(f'listbox {4}').Update(Classe)
                window.find_element('combo').Update(Categoria[0])
                window.refresh()                            
                window['coluna'].contents_changed() 
                

            except ValueError as erro:
                print("Coloque os valores corretamente")


        if event == 'Deletar':
                if Nome:
                    teste = sg.popup_ok_cancel('Deseja Mesmo Deletar este item?')
                    if teste == 'Cancel':
                        print('okay')
                    else:
                        y = values[f'listbox {1}'][0]
                        x = (y[0])
                        delete(x)
                        Nome = read_task()
                        rows = len(Nome)
                        Quantidade = read_task1()
                        Preco = read_task2()
                        Id = read_task3()
                        Classe = read_task4()
                        for i in range(cols):
                            listbox = window[f'listbox {i}']
                            listbox.Widget.configure(height=rows)  
                        window.find_element(f'listbox {0}').Update(Id)
                        window.find_element(f'listbox {1}').Update(Nome)
                        window.find_element(f'listbox {2}').Update(Quantidade)
                        window.find_element(f'listbox {3}').Update(Preco)
                        window.find_element(f'listbox {4}').Update(Classe)
                        window.refresh()                            
                        window['coluna'].contents_changed() 

        if event == '➕Adicionar classe➕':
            teste = sg.popup_get_text('Digite o nome da classe:', 'Ex: Capacitor')
            Categoria = str(teste)
            
            if Categoria == 'None':
                print('okay')
            elif not Categoria:
                print('okay')
            
            elif Categoria != '' or None:
                escreve_class(Categoria)
               
        

        if event == sg.WIN_CLOSED or event == 'Sair':
            window.close()
            break

        if event == '⬅ Voltar':
            window.close()
            initi()
            
            
            
        
####################################venda###################################################

with open("logo.ico", "rb") as f:
    my_icon = base64.b64encode(f.read())
    sg.set_options(icon=my_icon)


def JanelaEditId():
    sg.change_look_and_feel('DarkGreen')
    elayout = [
        [sg.Text('Digite o novo Id:'),sg.Input(key='edit_item')],
        [sg.Button('Enviar')]
              ]
    window = sg.Window('Novo Id',elayout)
    event, values = window.read()
    global edit
    edit = values['edit_item']
    window.close()

    if event == sg.WIN_CLOSED:
        window.close()

def JanelaEditPreco():
    sg.change_look_and_feel('DarkGreen')
    elayout = [
        [sg.Text('Digite o novo Preço:'),sg.Input(key='edit_item')],
        [sg.Button('Enviar')]
              ]
    window = sg.Window('Novo Preço',elayout)
    event, values = window.read()
    global edit
    edit = values['edit_item']
    window.close()

    if event == sg.WIN_CLOSED:
        window.close()
        

        

def JanelaVenda():
    sg.change_look_and_feel('DarkGreen')


    def yscroll(self, event):   # tkinter code
        if self.canvas.yview() == (0.0, 1.0):
            return
        if event.num == 5 or event.delta < 0:
            self.canvas.yview_scroll(1, "unit")
        elif event.num == 4 or event.delta > 0:
            self.canvas.yview_scroll(-1, "unit")
    
    
    sg.TkScrollableFrame.yscroll = yscroll

    Nome = '' 
    Quantidade = ''
    Preco = ''
    Id = ''
    Classe = ''

    cols = 5
    rows = 14
    col_width = 30
    
    
    all_listbox = [[sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {0}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(50, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {1}',
   select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {2}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {3}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE),sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {4}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE)]]

    


    vlayout = [
        [sg.Text('Digite o item que deseja fazer a venda:'), sg.Input(do_not_clear=False, size=(20,1),key='vender')],
        [sg.Text('ID',pad=((12,240),12)),sg.Text('Produto',pad=((12,155),12)),sg.Text('Quantidade',pad=((0,60),12)),sg.Text('Preço',pad=((0,55),12)),sg.Text('Classe')],
        [sg.Column(all_listbox, size=(811, 250), pad=(0, 0),scrollable=True,
         vertical_scroll_only=True, key='coluna')],
        [sg.Button('Consultar',size=(10, 1), font=(6),pad=((0,80),12)), sg.Button('Realizar venda'),sg.Button('Deletar',pad=((0,30),12)),sg.Button('Editar Preço'),sg.Button('Add estoque'),sg.Button('Editar Identifição')],
        [sg.Button('Sair'),sg.Button('⬅ Voltar')]
    ]

    window = sg.Window('Controle de estoque',vlayout)
    while True:
        event,values = window.read()
        
        try:
            try:
                if event.startswith('listbox'):
                        row = window[event].get_indexes()[0]
                        user_event = False
                        for i in range(cols):
                            listbox = window[f'listbox {i}']
                            listbox.set_value([])
                            listbox.Widget.selection_set(row)   # tkinter code
            except AttributeError:
                window.close()
        except IndexError:
            [sg.popup_ok('Primeiro Digite um Produto', font = (12))]
    
        if event == 'Consultar':
            c = values['vender']
            Id = busca0(c)
            Nome = busca1(c)
            rows = len(Nome)
            if rows < 11:
                rows = 14
            Quantidade = busca2(c)
            Preco = busca3(c)
            Classe = busca4(c)

            for i in range(cols):
                listbox = window[f'listbox {i}']
                listbox.Widget.configure(height=rows)  
            

            window.find_element(f'listbox {0}').Update(Id)
            window.find_element(f'listbox {1}').Update(Nome)
            window.find_element(f'listbox {2}').Update(Quantidade)
            window.find_element(f'listbox {3}').Update(Preco)
            window.find_element(f'listbox {4}').Update(Classe)
            
            window.refresh()                            
            window['coluna'].contents_changed() 
        try:
            
            if event == 'Realizar venda':
                if Nome:
                    x = values[f'listbox {1}'][0]
                    y = (x[0])
                    print(y)
                    Avenda()
                    QuantiTrat = int(Quanti)
                    print(QuantiTrat)
                    Quantidade = busca2_teste(y)
                    Vquantidade0 = (Quantidade[0])
                    Vquantidade = (Vquantidade0[0])
                    subtration = Vquantidade - QuantiTrat
                    q = subtration
                    h = y
                    vender(q, h)
                    Quantidade = busca2(c)
                    window.find_element(f'listbox {0}').Update(Id)
                    window.find_element(f'listbox {1}').Update(Nome)
                    window.find_element(f'listbox {2}').Update(Quantidade)
                    window.find_element(f'listbox {3}').Update(Preco)
                    window.find_element(f'listbox {4}').Update(Classe)
            

                    
            if event == 'Add estoque':
                if Nome:
                    x = values[f'listbox {1}'][0]
                    y = (x[0])
                    Add_Quanti()
                    AddQuantiTrat = int(AddQuanti)
                    AddQuantidade = busca2_teste(y)
                    Vquantidade1 = (AddQuantidade[0])
                    Vquantidade2 = (Vquantidade1[0])
                    addition = Vquantidade2 + AddQuantiTrat
                    q = addition
                    h = y
                    add_estoque(q, h)
                    Quantidade = busca2(c)
                    window.find_element(f'listbox {0}').Update(Id)
                    window.find_element(f'listbox {1}').Update(Nome)
                    window.find_element(f'listbox {2}').Update(Quantidade)
                    window.find_element(f'listbox {3}').Update(Preco)
                    window.find_element(f'listbox {4}').Update(Classe)
        except (TypeError, IndexError, ValueError):
            [sg.popup_ok('Selecione e Digite um valor valido', font = (12))]
        
                


        try:
            if event == 'Deletar':
                if Nome:
                    teste = sg.popup_ok_cancel('Deseja Mesmo Deletar este item?')
                    if teste == 'Cancel':
                        print('okay')
                    else:
                        y = values[f'listbox {1}'][0]
                        x = (y[0])
                        delete(x)
                        Id = busca0(c)
                        Nome = busca1(c)
                        Quantidade = busca2(c)
                        Preco = busca3(c)
                        Classe = busca4(c)
                        window.find_element(f'listbox {0}').Update(Id)
                        window.find_element(f'listbox {1}').Update(Nome)
                        window.find_element(f'listbox {2}').Update(Quantidade)
                        window.find_element(f'listbox {3}').Update(Preco)
                        window.find_element(f'listbox {4}').Update(Classe)
            
            if event == 'Editar Preço':
                if Nome:
                    x = values[f'listbox {1}'][0]
                    y = (x[0])
                    JanelaEditPreco()
                    p = edit
                    h = y
                    if not p:
                        [sg.popup_ok('Digite algum valor', font=(9))]
                        r = Preco[0]
                        p = r[0]
                    Editar(p, h)
                    Preco = busca3(c)
                    window.find_element(f'listbox {0}').Update(Id)
                    window.find_element(f'listbox {1}').Update(Nome)
                    window.find_element(f'listbox {2}').Update(Quantidade)
                    window.find_element(f'listbox {3}').Update(Preco)
                    window.find_element(f'listbox {4}').Update(Classe)

            if event == 'Editar Identifição':
                if Nome:
                    x = values[f'listbox {1}'][0]
                    y = (x[0])
                    JanelaEditId()
                    p = edit
                    h = y
                    if not p:
                        [sg.popup_ok('Digite algum valor', font=(9))]
                        r = Id[0]
                        p = r[0]
                    EditarID(p, h)
                    Id = busca0(c)
                    window.find_element(f'listbox {0}').Update(Id)
                    window.find_element(f'listbox {1}').Update(Nome)
                    window.find_element(f'listbox {2}').Update(Quantidade)
                    window.find_element(f'listbox {3}').Update(Preco)
                    window.find_element(f'listbox {4}').Update(Classe)
        except IndexError:
            [sg.popup_ok('Selecione e Digite um valor valido', font = (12))]        


        if event == sg.WIN_CLOSED or event == 'Sair':
            window.close()
            break
        if event == '⬅ Voltar':
            window.close()
            initi()
        
def Avenda():
    sg.change_look_and_feel('DarkGreen')
    #layout
    tlayout = [
        [sg.Text('Digite a quantidade:',size=(15,0)),sg.Input(size=(20,0), key='InQuanti')],
        [sg.Button('Vender')]
        
        ]
    #janela
    window = sg.Window("Controle de estoque", tlayout)


    event,values = window.read()
    global Quanti
    Quanti = values['InQuanti']
    if event == 'Vender':
        window.close()

def Add_Quanti():
    sg.change_look_and_feel('DarkGreen')
    #layout
    ylayout = [
        [sg.Text('Digite a quantidade:',size=(15,0)),sg.Input(size=(20,0), key='Adicionar_Quanti')],
        [sg.Button('Adicionar')]
        
        ]
    #janela
    window = sg.Window("Controle de estoque", ylayout)


    event,values = window.read()
    global AddQuanti
    AddQuanti = values['Adicionar_Quanti']
    if event == 'Adicionar':
        window.close()
    if event == sg.WIN_CLOSED:
        window.close()
        
    
    
#########################################filtragem##############################################


def Filtrar():
    with open("logo.ico", "rb") as f:
        my_icon = base64.b64encode(f.read())
        sg.set_options(icon=my_icon)       

    #mudar tema
    sg.change_look_and_feel('DarkGreen')



    def yscroll(self, event):   # tkinter code
        if self.canvas.yview() == (0.0, 1.0):
            return
        if event.num == 5 or event.delta < 0:
            self.canvas.yview_scroll(1, "unit")
        elif event.num == 4 or event.delta > 0:
            self.canvas.yview_scroll(-1, "unit")
    
    
    
    sg.TkScrollableFrame.yscroll = yscroll

    Nome = ''
    Quantidade = ''
    Preco = ''
    Id = ''
    Categoria = read_task5()

    
    cols = 4
    rows = 14
    col_width = 22
    
    
    all_listbox = [[sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {0}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(50, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {1}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {2}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE), sg.Listbox(Nome, size=(15, rows), pad=(2, 0),
    no_scrollbar=True, enable_events=True, key=f'listbox {3}',
    select_mode=sg.LISTBOX_SELECT_MODE_SINGLE)]]



    #layout
    layout = [
        [sg.Text('Selecione a classe que deseja ver:',size=(15,0)),sg.InputCombo((Categoria),size=(20,0),pad=((3,110),12),key='combo'),sg.Text('Ou digite a marca:',size=(15,0)), sg.Input(do_not_clear=False, size=(20,1),key='marca')],
        [sg.Button('Consultar')], 
         [sg.Text('Id'.center(col_width), pad=((5,140),12)),
          sg.Text('Nome'.center(col_width), pad=((5,110),12)),
          sg.Text('Quantidade'.center(col_width), pad=((0,0),11)),
          sg.Text('Preço'.center(col_width), pad=(0, 0))],
        [sg.Column(all_listbox, size=(700, 250), pad=(0, 0),scrollable=True,
         vertical_scroll_only=True, key='Coluna')],
        
        [sg.Button('Deletar',pad=((5,450),12)),sg.Button('Gerar Relatorio'),sg.Button('Abrir Relatorio')],
        [sg.Button('Sair'),sg.Button('⬅ Voltar',pad=((5,400),12)), sg.Button('Gerar Cotação'),sg.Button('Abrir Cotação')]
        ]
        
    #janela
    window = sg.Window("adicionar ao Estoque",layout)



    while True:
        event,values = window.read()

        try:
            try:
                if event.startswith('listbox'):
                    row = window[event].get_indexes()[0]
                    user_event = False
                    for i in range(cols):
                        listbox = window[f'listbox {i}']
                        listbox.set_value([])
                        listbox.Widget.selection_set(row)   # tkinter code
            except AttributeError:
                window.close()
            
            if event == 'Consultar':
                classe_nao_tratada = values['combo']
                f = str(classe_nao_tratada[0])
                Marca = values['marca']
                
                if not Marca: 
                    Id = filtrar2(f)
                    Nome = filtrar(f)
                    rows = len(Nome)
                    if rows < 13:
                        rows = 14
                    if not rows:
                        sg.popup_ok('Não ha nenhum produto dessa classe', font=(6))
                        rows = 10
                    Quantidade = filtrar3(f)
                    Preco = filtrar4(f)
                else:
                    Id = filtrar2_com_marca(Marca)
                    Nome = filtrar_com_marca(Marca)
                    Quantidade = filtrar3_com_marca(Marca)
                    Preco = filtrar4_com_marca(Marca)
                
                for i in range(cols):
                    listbox = window[f'listbox {i}']
                    listbox.Widget.configure(height=rows)  
                
                window.find_element(f'listbox {1}').Update(Nome)
                window.find_element(f'listbox {0}').Update(Id)
                window.find_element(f'listbox {2}').Update(Quantidade)
                window.find_element(f'listbox {3}').Update(Preco)
                window.refresh() 
                window['Coluna'].contents_changed()
                
                
                

        except IndexError:
            [
            [sg.popup_ok('Selecione alguma classe')]
            ]
        
      





    
        if event == 'Deletar':
            if Nome:
                teste = sg.popup_ok_cancel('Deseja Mesmo Deletar este item?')
                if teste == 'Cancel':
                        print('okay')
                else:
                    y = values[f'listbox {1}'][0]
                    x = (y[0])
                    delete(x)
                    Id = filtrar2(f)
                    Nome = filtrar(f)
                    Quantidade = filtrar3(f)
                    Preco = filtrar4(f)
                    window.find_element(f'listbox {1}').Update(Nome)
                    window.find_element(f'listbox {0}').Update(Id)
                    window.find_element(f'listbox {2}').Update(Quantidade)
                    window.find_element(f'listbox {3}').Update(Preco)
        try:
            if event == "Gerar Relatorio":
                lista = []

                
                for addname, addamount in zip(Nome, Quantidade):
                    lista.append([addname,addamount])
                
                
                doc = docx.Document()
                doc.add_heading('Relatorio' , 0)

                mTable = doc.add_table(rows=1, cols=3)
                mTable.style =  'Table Grid'
                cabeca = mTable.rows[0].cells
                cabeca[0].text = 'Id'
                cabeca[1].text = 'Nome'
                cabeca[2].text = 'Quantidade'

                
                for nome,quantidade in lista:
                    quantidade = str(quantidade)
                    lixo = "',()"
                    for i in range(0, len(lixo)):
                        quantidade = quantidade.replace(lixo[i], "")
                    row_cells = mTable.add_row().cells
                    row_cells[1].text = nome
                    row_cells[2].text = quantidade
                    
                  


                    
                doc.save('Relatorio.docx')
        except PermissionError:
                [
        [sg.popup_ok('Feche o arquivo caso esteja aberto, para gerar outro')]
            ]

        try:
    
    
            
            if event == "Gerar Cotação":
                
                for i in range(3000):
                    sg.PopupAnimated(sg.DEFAULT_BASE64_LOADING_GIF, background_color='green', transparent_color='green', time_between_frames=100)

                 

                arquivos = os.listdir(os.getcwd())
                for i in arquivos:
                #ler os arquivos que tem no diretorio e ver se existe cotacao.docx
                    if i == 'Cotacao.docx': 
                        verify = "achei"
                        break
                    else:
                        verify = "nao achei"

                        #aqui ele atribui um valor de nao existe caso nao encontre cotacao.docx
                
                if verify == "achei":
                    #se existir ele pega os dados que ja tinham reaplica e adiciona os novos
                        document = docx.Document('Cotacao.docx')
                        table = document.tables[0]
                        
                        data = []

                        keys = None
                        for i, row in enumerate(table.rows):
                            text = (cell.text for cell in row.cells)
                            if i == 0:
                                keys = tuple(text)
                                continue

                            row_data = dict(zip(keys, text))
                           
                            data.append(row_data)
   
                        
                        
                        lista = []

                        
                        for addname, addamount in zip(Nome, Quantidade):
                            lista.append([addname,addamount])
                        
                    
                        doc = docx.Document()
                        doc.add_heading('Cotação' , 0)
                        mTable = doc.add_table(rows=1, cols=4)
                        mTable.style =  'Table Grid'
                        cabeca = mTable.rows[0].cells
                        cabeca[0].text = 'Codigo'
                        cabeca[1].text = 'Produto'
                        cabeca[2].text = 'Quantidade'
                        cabeca[3].text = 'Preço'
                        
                        for name in data: 
                            #neste for ele readiciona os que ja existia
                            name = str(name["Produto"])
                            lixo = "',()[]"
                            for i in range(0, len(lixo)):
                                name = name.replace(lixo[i], "")
                            row_cells = mTable.add_row().cells
                            row_cells[1].text = name
                        
                        for nome in lista: #neste for ele adiciona os novos
                            nome = nome[0]
                            nome = str(nome)
                            lixo = "',()[]"
                            for i in range(0, len(lixo)):
                                nome = nome.replace(lixo[i], "")
                            row_cells = mTable.add_row().cells
                            row_cells[1].text = nome
                        doc.save('Cotacao.docx')
                        sg.PopupAnimated(None)
                elif verify == "nao achei":
                    #coloquei separado pois se eu colocasse dentro do for ele apagar os dados antigos
                    lista = []

                    
                    for addname, addamount in zip(Nome, Quantidade):
                        lista.append([addname,addamount])
                    
                    
                    doc = docx.Document()
                    doc.add_heading('Cotação' , 0)
                    mTable = doc.add_table(rows=1, cols=4)
                    mTable.style =  'Table Grid'
                    cabeca = mTable.rows[0].cells
                    cabeca[0].text = 'Codigo'
                    cabeca[1].text = 'Produto'
                    cabeca[2].text = 'Quantidade'
                    cabeca[3].text = 'Preço'
                    
                    
                    
                    for nome in lista: 
                        #aqui ele cria os dados novos
                        nome = nome[0]
                        nome = str(nome)
                        lixo = "',()[]"
                        for i in range(0, len(lixo)):
                            nome = nome.replace(lixo[i], "")
                        row_cells = mTable.add_row().cells
                        row_cells[1].text = nome
                  
                    doc.save('Cotacao.docx')
                    sg.PopupAnimated(None)

    
                  


                    
                
        except PermissionError:
                [
        [sg.popup_ok('Feche o arquivo caso esteja aberto, para gerar outro')]
            ]
        

        if event == "Abrir Relatorio":
            os.system("start Relatorio.docx")
        
        if event == "Abrir Cotação":
            os.system("start Cotacao.docx")
        
      
        if event == sg.WIN_CLOSED or event == 'Sair':
            window.close()
            break
             
        if event == '⬅ Voltar':
            window.close()
            initi()

########################################################################################################################################

def initi():
    with open("logo.ico", "rb") as f:
        my_icon = base64.b64encode(f.read())
        sg.set_options(icon=my_icon)



    sg.change_look_and_feel('DarkGreen')

    ilayout = [
        [sg.Text('Olá seja Bem-vindo, Oque deseja fazer?')],
        [sg.Button('Adicionar ao Estoque')],
        [sg.Button('Realizar uma venda')],
        [sg.Button('Fazer uma filtragem')],
        [sg.Button('Adicionar Tvs')],
        [sg.Button('Sair')],
        
    ]

    window = sg.Window('Controle de estoque',ilayout)

    button, values = window.read()
    if button == 'Adicionar ao Estoque':
        window.close()
        AdicionarItem()

    


    if button == 'Realizar uma venda':
        window.close()
        JanelaVenda()
    if button == 'Fazer uma filtragem':
        window.close()
        Filtrar()
    if button == 'Sair':
        window.close()
    
    if button == 'Adicionar Tvs':
        window.close()
        AdicionarTelevisao()



initi()